import re
import io
from typing import Optional
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract
from docx import Document as DocxDocument


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfminer (better) with PyPDF2 fallback."""
    try:
        text = pdfminer_extract(io.BytesIO(file_bytes))
        if text and len(text.strip()) > 50:
            return text
    except Exception:
        pass

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as e:
        raise ValueError(f"Could not extract text from PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Could not extract text from DOCX: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to correct extractor based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─── Regex-based field extractors ───────────────────────────────────────────

def extract_email(text: str) -> Optional[str]:
    pattern = r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    pattern = r"(\+?\d{1,3}[\s\-]?)?(\(?\d{3}\)?[\s\-]?)(\d{3}[\s\-]?\d{4})"
    match = re.search(pattern, text)
    return match.group(0).strip() if match else None


def extract_name(text: str) -> Optional[str]:
    """Heuristic: first non-empty line that looks like a name."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:5]:
        # Skip lines with email/phone/url
        if any(c in line for c in ["@", "http", "linkedin", "github", "+"]):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
            return line
    return lines[0] if lines else None


def extract_location(text: str) -> Optional[str]:
    patterns = [
        r"(?:Location|Address|City)[:\s]+([^\n]+)",
        r"\b([A-Z][a-z]+(?:,\s*[A-Z]{2})?(?:,\s*[A-Z][a-z]+)?)\b",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(1).strip()
    return None


# Common skill keywords
TECHNICAL_SKILLS_KEYWORDS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby",
    "react", "angular", "vue", "node.js", "fastapi", "django", "flask", "spring",
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins",
    "git", "linux", "bash", "html", "css", "graphql", "rest", "api",
    "machine learning", "deep learning", "tensorflow", "pytorch", "scikit-learn",
    "pandas", "numpy", "spark", "hadoop", "kafka",
    "figma", "photoshop", "tableau", "power bi", "excel",
    "r", "scala", "swift", "kotlin", "php", "laravel"
]

SOFT_SKILLS_KEYWORDS = [
    "communication", "teamwork", "leadership", "problem solving", "critical thinking",
    "time management", "adaptability", "creativity", "collaboration", "attention to detail",
    "project management", "mentoring", "analytical", "presentation", "negotiation",
    "decision making", "conflict resolution", "emotional intelligence", "initiative"
]


def extract_skills(text: str):
    text_lower = text.lower()
    technical = [s for s in TECHNICAL_SKILLS_KEYWORDS if s in text_lower]
    soft = [s for s in SOFT_SKILLS_KEYWORDS if s in text_lower]
    return technical, soft


def extract_experience_years(text: str) -> Optional[float]:
    patterns = [
        r"(\d+)\+?\s*years?\s+(?:of\s+)?experience",
        r"experience[:\s]+(\d+)\+?\s*years?",
    ]
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            return float(m.group(1))
    return None


def extract_education(text: str):
    degrees = ["bachelor", "master", "phd", "doctorate", "b.tech", "m.tech", "b.e", "m.e",
               "bsc", "msc", "mba", "b.com", "b.a", "m.a", "associate"]
    lines = text.lower().split("\n")
    education = []
    for i, line in enumerate(lines):
        if any(deg in line for deg in degrees):
            education.append({"degree": lines[i].strip(), "details": lines[i + 1].strip() if i + 1 < len(lines) else ""})
    return education[:5]


def extract_certifications(text: str):
    certs = []
    patterns = [
        r"(?:certified|certification|certificate)[^.\n]*[.\n]",
        r"AWS\s+\w+",
        r"Google\s+\w+\s+\w+",
        r"Microsoft\s+\w+",
        r"PMP|CISSP|CPA|CFA|CCNA|CCNP|CISA",
    ]
    for p in patterns:
        for m in re.finditer(p, text, re.IGNORECASE):
            cert = m.group(0).strip().rstrip(".")
            if cert not in certs:
                certs.append(cert)
    return certs[:10]


def extract_languages(text: str):
    languages = ["english", "hindi", "spanish", "french", "german", "mandarin",
                 "arabic", "portuguese", "japanese", "korean", "italian", "russian",
                 "urdu", "bengali", "tamil", "telugu", "marathi", "gujarati"]
    text_lower = text.lower()
    return [lang.capitalize() for lang in languages if lang in text_lower]


def parse_resume(text: str) -> dict:
    """Parse all fields from resume text."""
    technical_skills, soft_skills = extract_skills(text)
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "location": extract_location(text),
        "summary": None,
        "skills": list(set(technical_skills + soft_skills)),
        "technical_skills": technical_skills,
        "soft_skills": soft_skills,
        "experience": [],
        "total_experience_years": extract_experience_years(text),
        "education": extract_education(text),
        "certifications": extract_certifications(text),
        "projects": [],
        "languages": extract_languages(text),
    }
